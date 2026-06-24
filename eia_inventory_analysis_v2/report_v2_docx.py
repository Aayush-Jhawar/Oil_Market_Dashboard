"""
================================================================================
  EIA INVENTORY-IMPACT  v2  -  WORD REPORTS (two .docx)
================================================================================
Produces TWO documents from one pipeline run:
  1. EIA_Inventory_Deliverable_v2.docx  - concise, structured to the four asks
     (expectation / products-spreads / top-3 factors / framework) + focus areas.
  2. EIA_Inventory_Detailed_v2.docx      - full exploratory research, with a
     glossary, plain-English table captions, all figures, and the four focus
     areas in depth.

Exploratory research framing throughout.
"""
import os
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import report_v2 as r2

OUT = r2.OUT
DELIV = os.path.join(OUT, "EIA_Inventory_Deliverable_v2.docx")
DETAIL = os.path.join(OUT, "EIA_Inventory_Detailed_v2.docx")

# human-readable column names so tables don't need decoding
COLMAP = {
    "release_date": "Release date", "signal": "Signal tested", "n": "N",
    "EIA_Crude__actual": "EIA crude actual (Mbbl)", "EIA_Crude__consensus": "Consensus (Mbbl)",
    "API_Crude__actual": "API actual (Mbbl)", "crude_surp_cons": "Surprise vs consensus (Mbbl)",
    "crude_surp_API": "EIA minus API (Mbbl)", "WTI_post10_pct": "WTI +10min (%)",
    "WTI_spike_pct": "WTI peak move (%)", "hit_at_release": "Hit @release",
    "hit_+5min": "Hit +5min", "hit_+10min": "Hit +10min", "corr_+10min": "Corr (surprise,+10min)",
    "bucket": "Condition", "instrument": "Instrument", "hit": "Hit rate", "corr": "Corr",
    "hit_rate": "Hit rate", "mean_spike_pct": "Mean peak move (%)",
    "mean_abs_surprise": "Mean |surprise| (Mbbl)", "mean_rvol": "Mean realised vol",
    "api_divergence_rate": "API/EIA disagree rate", "EIA_hit_rate": "EIA-signal hit rate",
    "mean_WTI_spike_pct": "Mean WTI peak move (%)", "agree": "API vs EIA",
    "quarter": "Quarter", "season": "Season", "vol_regime": "Vol regime", "grp": "Move-size group",
    "acc_at_release": "Accuracy at release", "acc_+5min": "Accuracy +5min",
    "acc_+10min": "Accuracy +10min", "n_trades": "Trades (n)", "minute": "Minute from print",
    "WTI_absmove": "WTI mean |move| (%)", "WTI_step": "WTI per-minute step (%)",
    "WTI_share_of_10min": "Share of the +10min move", "Brent_absmove": "Brent mean |move| (%)",
    "HO_absmove": "HO mean |move| (%)",
}


def add_table(doc, df, caption=None, floatfmt=3, max_rows=None, index_label=None):
    d = df.copy()
    if index_label is not None:
        d = d.reset_index().rename(columns={"index": index_label})
    if max_rows:
        d = d.head(max_rows)
    cols = [COLMAP.get(c, c) for c in d.columns]
    if caption:
        cp = doc.add_paragraph(); run = cp.add_run(caption); run.italic = True; run.font.size = Pt(9)
    t = doc.add_table(rows=1, cols=len(cols)); t.style = "Light Grid Accent 1"
    for j, c in enumerate(cols):
        t.rows[0].cells[j].paragraphs[0].add_run(str(c)).bold = True
    is_float = {c: pd.api.types.is_float_dtype(d[c]) for c in d.columns}
    is_int = {c: pd.api.types.is_integer_dtype(d[c]) for c in d.columns}
    for _, row in d.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(d.columns):
            v = row[c]
            if pd.isna(v):
                cells[j].text = ""
            elif is_float[c]:
                cells[j].text = f"{v:.{floatfmt}f}"
            elif is_int[c]:
                cells[j].text = str(int(v))
            else:
                cells[j].text = str(v)
    doc.add_paragraph()


def H(doc, t, lvl=1):
    doc.add_heading(t, level=lvl)


# ----------------------------------------------------------- the call ---------
def make_call(P):
    df = P["df"]; edges = P["edges"]
    d = df.dropna(subset=["crude_surp_cons"]).sort_values("release_date")
    latest = d.iloc[-1]
    sc = latest["crude_surp_cons"]
    qlo, qhi = df["crude_surp_cons"].abs().quantile([1 / 3, 2 / 3])
    size = "large" if abs(sc) > qhi else ("small" if abs(sc) < qlo else "mid")
    fund = "bullish" if sc < 0 else ("bearish" if sc > 0 else "neutral")
    if size == "large":
        view = (f"Neutral, lean to fade. The number reads {fund} on fundamentals, "
                "but large surprises like this one have tended to reverse")
    elif size == "small":
        view = f"{fund.capitalize()}, low conviction. Small surprises tend to follow through"
    else:
        view = f"{fund.capitalize()}, weak signal and low conviction"

    def gh(b, inst="WTI +10min"):
        v = edges[(edges.bucket == b) & (edges.instrument == inst)]["hit"].values
        return v[0] if len(v) else np.nan
    return {"latest": latest, "sc": sc, "size": size, "fund": fund, "view": view,
            "small_hit": gh("|surp| small"), "large_hit": gh("|surp| large"),
            "api_div": latest.get("crude_surp_API", np.nan)}


# ----------------------------------------------------------- doc 1: deliverable
def build_deliverable(P, figs, call):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)
    t = doc.add_heading("EIA Crude Inventory: Market Impact Deliverable", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph("Exploratory research. Consensus and API surprise versus intraday "
                          f"price, across {len(P['df'])} releases from 2021 to 2026.")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER; s.runs[0].italic = True
    doc.add_paragraph()

    lat = call["latest"]
    H(doc, "1. Directional expectation (most recent release)", 1)
    doc.add_paragraph(
        f"Take the most recent print ({lat['release_date']}) as a worked example. EIA crude "
        f"came in at {lat['EIA_Crude__actual']:+.1f} Mbbl against a consensus of "
        f"{lat['EIA_Crude__consensus']:+.1f}, so the surprise was {call['sc']:+.1f} Mbbl, "
        f"which counts as a {call['size']} surprise.")
    p = doc.add_paragraph(); p.add_run("View: ").bold = True
    p.add_run(call["view"] + ". Confidence is low.")
    doc.add_paragraph(
        "Best structure: almost all of the (small) inventory signal sits in the M1 outright, "
        "so calendar spreads and cracks add very little. The only expression that has paid "
        "consistently is fading outsized prints in WTI M1, where large surprises reverse "
        f"(hit rate {call['large_hit']:.2f}). Small surprises tend to follow through, but only "
        f"weakly (hit rate {call['small_hit']:.2f}).", style="List Bullet")
    doc.add_picture(figs["conditioned"], width=Inches(6.2))

    H(doc, "2. Products and spreads most likely to be affected", 1)
    doc.add_paragraph("Brent (ICE) shows the largest and fastest reaction, and prices the crude "
                      "number correctly at the moment of release more often than WTI does.", style="List Bullet")
    doc.add_paragraph("WTI (NYMEX) is the most crude-specific instrument, with the cleanest link "
                      "to the US surprise, so it is the natural home for the fade-the-extreme trade.", style="List Bullet")
    doc.add_paragraph("Heating oil is the one to watch for the distillate surprise. The "
                      "WTI/Brent and heating-oil crack spreads add little over the outright.", style="List Bullet")
    doc.add_paragraph("Gasoline: EIA inventory data exists, but we did not have RBOB one-minute "
                      "prices, so it stays inventory-only for now (flagged for a future data pull).", style="List Bullet")

    H(doc, "3. Top three factors driving the view", 1)
    doc.add_paragraph("Surprise size versus consensus. This is the switch that matters most: "
                      "small surprises follow the print, large ones fade it. It is the most "
                      "actionable thing we found.", style="List Bullet")
    doc.add_paragraph("Volatility regime. This drives how big the reaction is, and whether the "
                      "release is worth trading at all. High-volatility weeks move about 60% more.", style="List Bullet")
    doc.add_paragraph("The Tuesday API read. When the Tuesday API and the Wednesday EIA agree, "
                      "the EIA reaction is cleaner. When they diverge, the Tuesday move tends to "
                      "pre-empt and muddy the Wednesday signal.", style="List Bullet")

    H(doc, "4. Framework in brief", 1)
    doc.add_paragraph(
        "On Tuesday afternoon the API estimate gives a first read of the surprise against "
        "consensus. On Wednesday the EIA print either confirms it or pulls the other way. We "
        "measure the surprise as actual minus consensus, and separately as EIA minus API, then "
        "trade the magnitude rather than the direction: size the position by the surprise and "
        "the volatility regime, lean with small surprises, fade the extreme prints, and keep in "
        "mind that inventories are only one input. On macro-heavy days they are not the "
        "marginal driver.")

    H(doc, "Signal accuracy and execution timing", 1)
    doc.add_paragraph(
        "Following every print blindly is a coin flip (around 0.50). The rule we actually trade, "
        "follow small surprises and fade large ones, is right about 60% of the time on WTI, and "
        "fading a large print is right roughly two times in three at the release instant. The "
        "accuracy holds up across the +5 and +10 minute marks.")
    add_table(doc, P["accuracy"],
              caption="Accuracy of the generated signal on WTI, by horizon. 0.50 is a coin "
              "flip. 'Trades' is the number of releases the signal would have acted on.")
    doc.add_paragraph(
        "On timing, the move lands at the release minute itself. About 60% of the eventual "
        "10-minute move is in within two minutes and roughly 80% within five, so the window "
        "worth trading is the first few minutes. Brent reacts fastest of the three products.")

    H(doc, "Focus areas (summary)", 1)
    doc.add_paragraph("Historical reactions: direction is close to a coin flip even against the "
                      "real consensus, because the print is priced in within minutes.", style="List Bullet")
    doc.add_paragraph("When inventories mattered: the biggest moves track the volatility regime "
                      "rather than the surprise size, so high-volatility weeks are when it counts.", style="List Bullet")
    doc.add_paragraph("Seasonal and regime effects: seasonality is weak (Q3 is marginally the "
                      "best), and the volatility regime is the dominant conditioner.", style="List Bullet")
    doc.add_paragraph("Amplifiers and offsets: high volatility and agreement between the API and "
                      "EIA amplify a clean signal, while an outsized surprise or a competing macro "
                      "driver offsets or reverses it.", style="List Bullet")
    doc.add_paragraph()
    doc.add_paragraph("This is exploratory research: directional findings on about five years of "
                      "weekly data, not a production trading signal.").runs[0].italic = True
    doc.save(DELIV)
    print(f"  saved {DELIV}")


# ----------------------------------------------------------- doc 2: detailed --
GLOSSARY = [
    ("EIA release", "US Energy Information Administration weekly petroleum inventory report, out Wednesday at 10:30 ET (Thursday at 11:00 on holiday weeks)."),
    ("API estimate", "The American Petroleum Institute's private inventory estimate, released the afternoon before (Tuesday around 16:30 ET). It is an early read."),
    ("Consensus", "The median analyst forecast for the EIA number, in other words what the market has already priced in."),
    ("Surprise", "Actual minus expectation. We use the surprise against consensus and EIA minus API. A positive number is a bigger build than expected (bearish); a negative number is a bigger draw (bullish)."),
    ("M1 outright", "The front-month futures contract price (WTI is CL, Brent is LCO, heating oil is HO)."),
    ("Reaction window", "The move in the M1 price from the five-minute average before the release to: the release minute itself, the +5-minute average, and the +10-minute average."),
    ("Peak move (spike)", "The largest absolute percentage move from the pre-release level during the +10 minutes. It is a measure of magnitude, not direction."),
    ("Hit rate", "The share of releases where the surprise-implied direction (a draw points up, a build points down) matched the actual move. 0.50 is a coin flip."),
    ("Corr", "The Pearson correlation between the surprise and the price move. Fundamentals imply a negative correlation, since a bigger build should push price down."),
    ("Follow vs fade", "Follow means price moves in the surprise direction; fade means it reverses against it."),
    ("Realised vol (20d)", "The standard deviation of WTI daily returns over the prior 20 sessions. It stands in for the prevailing volatility regime."),
    ("WTI/Brent and HO crack", "Cross-product spreads (relative percentage moves) we tested as alternative ways to express the trade."),
    ("Signal accuracy", "How often the generated signal pointed the right way, measured on the releases the signal would actually have traded. We report the naive 'follow every print' version and the rule we trade (follow small surprises, fade large ones)."),
    ("Intraday reaction curve", "The average size of the move at each minute after the print, which shows the time range over which the bulk of the move happens."),
    ("Share of the +10min move", "At a given minute, the average move so far divided by the average move at +10 minutes. It reads as the fraction of the total reaction that is already in."),
]


def build_detailed(P, figs, call):
    df = P["df"]
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)
    t = doc.add_heading("EIA Crude Inventory Market Impact: Detailed Exploratory Report", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph("Consensus and API surprise versus intraday price reaction. Full "
                          "methodology, glossary, and figures.")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER; s.runs[0].italic = True
    doc.add_paragraph()

    H(doc, "How to read this report (glossary)", 1)
    doc.add_paragraph("Every term and table heading used below is defined here, so the tables "
                      "should not need any decoding.")
    gt = doc.add_table(rows=1, cols=2); gt.style = "Light Grid Accent 1"
    gt.rows[0].cells[0].paragraphs[0].add_run("Term").bold = True
    gt.rows[0].cells[1].paragraphs[0].add_run("Meaning").bold = True
    for term, mean in GLOSSARY:
        c = gt.add_row().cells; c[0].text = term; c[1].text = mean
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run("Sign convention. ").bold = True
    p.add_run("Throughout, a draw (inventory fell more than expected) is treated as bullish "
              "and a build as bearish. A release counts as a hit when the move agreed with "
              "that convention.")
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run("Data provenance and validity. ").bold = True
    p.add_run("Inventory actuals and consensus, plus the API estimate, were scraped from "
              "investing.com's economic calendar. Production and refinery utilisation come from "
              "EIA directly, and the one-minute prices come from the workspace futures files. We "
              "checked that the data is real and aligned: the workspace EIA changes match the "
              "official EIA series, and workspace WTI daily returns correlate 0.90 with real WTI.")
    doc.add_page_break()

    # ---- 1 ----
    H(doc, "1. Historical inventory releases and market reactions", 1)
    doc.add_paragraph(
        f"Across {len(df)} weekly EIA crude releases from 2021 to 2026, we line up the surprise "
        "(actual minus consensus) against the M1 move in the minutes after the 10:30 ET print. "
        "Figure 1 plots the surprise against the +10-minute move for WTI and Brent. If the "
        "print drove price on fundamentals, the cloud would slope downward, since a bigger "
        "build should mean a lower price. It barely does.")
    doc.add_picture(figs["scatter"], width=Inches(6.4))
    doc.add_picture(figs["hitrate"], width=Inches(6.2))
    add_table(doc, P["comp"],
              caption="Table 1. Directional hit rate by signal and horizon (at release, +5 and "
              "+10 min). 0.50 is a coin flip. 'Corr' is the surprise against the +10-minute "
              "move, where a negative value is the fundamentally correct sign.")
    doc.add_paragraph(
        "What it says: even using the correct expectation (consensus or API), the direction of "
        "the immediate move is close to a coin flip at every horizon. The crude print is "
        "largely priced in within minutes. Brent reacts a touch more correctly right at the "
        "release, while WTI is essentially random in the first minutes.")

    # ---- 2 intraday timing ----
    H(doc, "2. How fast the move happens (intraday timing)", 1)
    doc.add_paragraph(
        "Having seen that direction is close to random, the next question is timing: over what "
        "stretch of minutes does the reaction actually happen? We average the absolute move at "
        "each minute relative to the pre-release level, across all releases. Figure 9 shows the "
        "build-up; the table reads off how much of the eventual 10-minute move is in by each "
        "minute.")
    doc.add_picture(figs["curve"], width=Inches(6.4))
    ctbl = P["curve"][["minute", "WTI_absmove", "WTI_step", "WTI_share_of_10min",
                       "Brent_absmove", "HO_absmove"]].copy()
    ctbl = ctbl[(ctbl["minute"] >= 0) & (ctbl["minute"] <= 10)]
    ctbl["minute"] = ctbl["minute"].astype(int)
    add_table(doc, ctbl, floatfmt=3,
              caption="Table 2. Mean absolute move by minute from the print. 'Per-minute step' "
              "is the extra move added that minute; 'share of the +10min move' is how much of "
              "the reaction is already in. The pre-release minutes (not shown) sit at a noise "
              "floor of roughly 0.03 to 0.05%.")
    doc.add_paragraph(
        "The reaction lands at the release minute itself, which carries the single largest step. "
        "By two minutes after the print about 60% of the eventual 10-minute move is already in, "
        "and by five minutes roughly 80%. Brent moves fastest of the three, more than doubling "
        "its pre-release drift in the first minute. The practical read is that the window worth "
        "trading is the first few minutes; after that you are chasing a move that has mostly "
        "happened.")

    # ---- 3 signal accuracy ----
    H(doc, "3. Signal generation accuracy", 1)
    doc.add_paragraph(
        "This is the scorecard for the generated signal on WTI: how often it pointed the right "
        "way, measured only on the releases it would actually have traded. We show the naive "
        "'follow every print' version against the rule we trade, which follows small surprises "
        "and fades large ones.")
    doc.add_picture(figs["accuracy"], width=Inches(6.2))
    add_table(doc, P["accuracy"],
              caption="Table 3. Signal accuracy on WTI by horizon. 0.50 is a coin flip. "
              "'Trades' is the number of releases the signal acted on (the rule skips the middle "
              "tercile, so it trades fewer than the naive version).")
    doc.add_paragraph(
        "Following every print is a coin flip, as expected. The rule lifts accuracy to about "
        "0.60 and holds it across all three horizons. The work is done by the fade on large "
        "prints, which is right roughly 68% of the time at the release instant and about 62% by "
        "ten minutes; the follow on small prints adds a smaller, steadier edge near 0.57 to "
        "0.59. Accuracy is highest right at the release and decays slightly as the move matures, "
        "which fits the timing picture above.")

    # ---- 4 ----
    H(doc, "4. When inventories mattered, and when they didn't", 1)
    doc.add_paragraph(
        "We split the releases into terciles by the size of the WTI peak move (muted, mid, and "
        "mattered) and ask what sets the big-move weeks apart.")
    doc.add_picture(figs["when"], width=Inches(6.4))
    add_table(doc, P["when"], index_label="grp",
              caption="Table 4. Characteristics by size of the WTI move. The mean absolute "
              "surprise is roughly flat across the groups while realised vol rises, so the moves "
              "are driven by the volatility regime rather than by how big the surprise was.")
    doc.add_paragraph(
        "The takeaway: a release matters, in the sense of a big reaction, mainly when the market "
        "is already in a high-volatility regime, not simply because the surprise was large. "
        "Inventory days with a quiet tape tend to stay quiet whatever the number is.")

    # ---- 5 ----
    H(doc, "5. Seasonal and regime effects", 1)
    doc.add_picture(figs["seasonal"], width=Inches(6.4))
    add_table(doc, P["seasonal"]["quarter"], index_label="quarter",
              caption="Table 5a. Hit rate and mean peak move by calendar quarter.")
    if "vol_regime" in P["seasonal"]:
        add_table(doc, P["seasonal"]["vol_regime"], index_label="vol_regime",
                  caption="Table 5b. By prior realised-volatility regime (terciles).")
    doc.add_paragraph(
        "Seasonality is weak: Q3 is marginally the best for the signal, and the driving season "
        "looks about the same as the off-season. The clearer regime effect is volatility. In "
        "high-vol regimes the reaction is both bigger and slightly more directional.")

    # ---- 6 ----
    H(doc, "6. The conditioned edge and other drivers (what amplifies and what offsets)", 1)
    doc.add_paragraph(
        "Unconditionally the signal is a coin flip, so we went looking for an edge inside "
        "conditions (Figure 5). The pattern that holds up is simple: small surprises follow the "
        "print, large surprises fade it.")
    doc.add_picture(figs["conditioned"], width=Inches(6.4))
    add_table(doc, P["edges"],
              caption="Table 6. Conditioned hit rate and correlation by bucket and instrument. "
              "'Condition' is the subset. A hit below 0.50 with a positive correlation means the "
              "tape fades the surprise. The large-surprise WTI hit is about 0.39 (a fade), the "
              "small-surprise hit about 0.60 (a follow).")
    doc.add_paragraph(
        f"The fade on large prints (hit rate {call['large_hit']:.2f}) is the one tradeable, "
        "conditioned signal. An outsized number is usually already leaked by the Tuesday API "
        "and over-positioned into, so the Wednesday EIA confirmation tends to mean-revert.")
    H(doc, "The Tuesday API as an amplifier or an offset", 2)
    doc.add_picture(figs["api_eia"], width=Inches(6.2))
    add_table(doc, P["api_g"], index_label="agree",
              caption="Table 7. Outcomes when the Tuesday API and Wednesday EIA surprise the "
              "same way (agree) versus the opposite way (diverge).")
    doc.add_paragraph(
        "What amplifies a clean signal: a high-vol regime and agreement between the API and EIA. "
        "What offsets or muddies it: an outsized surprise (which fades), a Tuesday API that has "
        "already moved price, and competing macro or geopolitical drivers.")

    # ---- 7 magnitude ----
    H(doc, "7. Magnitude, the part you can condition on", 1)
    doc.add_picture(figs["magnitude"], width=Inches(6.4))
    doc.add_paragraph(
        "Direction is hard, but the size of the move scales with the surprise. That is useful "
        "for position sizing and for deciding whether a given release is worth trading at all.")

    # ---- 8 discrepancies ----
    H(doc, "8. Discrepancy investigation", 1)
    doc.add_paragraph(
        f"Of the {P['n_material']} material-surprise releases (absolute surprise above 0.5 "
        f"Mbbl), {len(P['disc'])} of them "
        f"({len(P['disc'])/max(P['n_material'],1)*100:.0f}%) saw the tape go against the "
        "surprise-implied direction. A sample is shown below.")
    add_table(doc, P["disc"], max_rows=12,
              caption="Table 8. Releases where the print and the tape disagreed. Large-surprise "
              "cases dominate (the fade effect), and the rest tend to coincide with a competing "
              "macro driver.")

    # ---- 9 framework / limits ----
    H(doc, "9. Framework and limitations", 1)
    doc.add_paragraph(
        "The framework, in one breath: read the Tuesday API, then the Wednesday EIA surprise "
        "against consensus; trade the magnitude (sizing by the surprise and the volatility "
        "regime); lean with small surprises; fade the extreme prints; prefer WTI M1 for the "
        "crude-specific expression and Brent for the fastest reaction; and treat inventories as "
        "one input among the macro drivers.")
    doc.add_paragraph("A few limitations to keep in mind. Consensus and API coverage depends on "
                      "investing.com, so a handful of weeks are missing. RBOB one-minute prices "
                      "were unavailable, so gasoline is inventory-only. The futures are "
                      "back-adjusted continuous series, so intraday percentage moves are exact "
                      "but dollar spreads are approximate. And with only about five years of "
                      "weekly data, this stays exploratory, with the usual multiple-testing "
                      "caveats.", style="List Bullet")
    doc.save(DETAIL)
    print(f"  saved {DETAIL}")


def main():
    print("=" * 60); print("  BUILDING v2 WORD REPORTS (deliverable + detailed)"); print("=" * 60)
    P = r2.run_pipeline()
    figs = r2.build_all_figures(P)
    call = make_call(P)
    build_deliverable(P, figs, call)
    build_detailed(P, figs, call)
    print("  DONE")


if __name__ == "__main__":
    main()
